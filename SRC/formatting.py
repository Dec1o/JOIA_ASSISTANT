from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx import Document
import datetime
import os

def criar_docx(conteudo, nome_arquivo_original, output_dir="OUTPUTS"):
    # Verifica se conteudo é uma string, caso contrário, converte
    if not isinstance(conteudo, str):
        conteudo = str(conteudo)

    # Remove caracteres inválidos do nome do arquivo original
    nome_arquivo_original = os.path.splitext(os.path.basename(nome_arquivo_original))[0]
    nome_arquivo_original = "".join(c for c in nome_arquivo_original if c.isalnum() or c in (' ', '-', '_')).strip()

    # Obtém a data e hora atual no formato desejado
    data_hora_atual = datetime.datetime.now().strftime("%d-%m-%Y-%H-%M")

    # Define o nome do arquivo com base no nome original
    nome_arquivo = f"ANALIZADO-{nome_arquivo_original}-{data_hora_atual}.docx"
    
    # Garante que o arquivo será salvo no diretório correto
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    caminho_completo = os.path.join(output_dir, nome_arquivo)

    # Cria um novo documento
    doc = Document()

    # Ajusta as margens do documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    # Adiciona cabeçalho e rodapé
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "Relatório de Análise"
    header_paragraph.style.font.size = Pt(10)
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = f"Documento analisado: {nome_arquivo_original} | Data: {data_hora_atual}"
    footer_paragraph.style.font.size = Pt(10)
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adiciona título principal
    heading = doc.add_heading('Resposta à Consulta', level=1)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # Sempre usa a cor preta

    # Adiciona a resposta ao documento com formatação
    p = doc.add_paragraph()
    run = p.add_run(conteudo)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)  # Sempre usa a cor preta
    run.bold = True
    p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph('Caso ainda tenha alguma dúvida, por favor entre em contato.')

    # Salva o documento com o nome personalizado
    doc.save(caminho_completo)

    return caminho_completo
